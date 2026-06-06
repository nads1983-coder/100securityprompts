from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BRAND = "100SecurityPrompts.com"
GOLD = RGBColor(214, 170, 67)
GOLD_DARK = RGBColor(122, 90, 0)
INK = RGBColor(20, 20, 20)
MUTED = RGBColor(88, 88, 88)
BLACK = RGBColor(7, 7, 7)

ROOT = Path("/Users/nadinepierre/Documents/New project")
SOURCES = [
    {
        "src": Path(
            "/Users/nadinepierre/Library/Mobile Documents/com~apple~CloudDocs/"
            "Security_Manager_AI_Toolkit_Commercial_Edition.docx"
        ),
        "docx": ROOT / "public/downloads/100-ai-prompts-for-security-managers.docx",
        "title": "100 AI Prompts for Security Managers",
        "subtitle": (
            "Commercial Edition | Practical ChatGPT prompts for security "
            "managers, supervisors and corporate security teams"
        ),
        "delivery": "Main Product PDF",
    },
    {
        "src": Path(
            "/Users/nadinepierre/Library/Mobile Documents/com~apple~CloudDocs/"
            "25_Real_World_Security_Manager_AI_Use_Cases.docx"
        ),
        "docx": ROOT
        / "public/downloads/25-real-world-security-manager-ai-use-cases-bonus.docx",
        "title": "25 Real-World Security Manager AI Use Cases",
        "subtitle": "Bonus resource for the 100 AI Prompts for Security Managers product",
        "delivery": "Bonus PDF",
    },
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_bottom_border(paragraph, color="D6AA43", size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_run(run, size=None, bold=None, color=None, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def apply_base_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    title.font.size = Pt(30)
    title.font.bold = True
    title.font.color.rgb = INK
    title.paragraph_format.space_after = Pt(6)

    for name, size, color, before, after in [
        ("Heading 1", 16, GOLD_DARK, 18, 10),
        ("Heading 2", 13, GOLD_DARK, 14, 7),
        ("Heading 3", 12, MUTED, 10, 5),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_footer(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(f"{BRAND} | Digital product for personal professional use")
    style_run(run, size=8.5, color=MUTED)


def add_cover(doc, title, subtitle, delivery):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    r = p.add_run(BRAND)
    style_run(r, size=15, bold=True, color=GOLD)
    add_bottom_border(p)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(72)

    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.paragraph_format.space_after = Pt(12)
    r = h.add_run(title)
    style_run(r, size=34, bold=True, color=INK)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(32)
    r = sub.add_run(subtitle)
    style_run(r, size=13, color=MUTED)

    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "070707")
    cell.text = ""
    badge = cell.paragraphs[0]
    badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
    badge.paragraph_format.space_before = Pt(18)
    badge.paragraph_format.space_after = Pt(18)
    r = badge.add_run(f"{delivery} | Instant download | One-time purchase")
    style_run(r, size=12, bold=True, color=GOLD)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(18)
    r = note.add_run(
        "Use these prompts as practical drafting support. Always verify facts, "
        "protect confidential information and follow site procedures."
    )
    style_run(r, size=10.5, color=MUTED)

    doc.add_page_break()


def add_intro(doc, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(BRAND)
    style_run(r, size=11, bold=True, color=GOLD)
    add_bottom_border(p, size="6")

    h = doc.add_paragraph(title, style="Title")
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    intro = doc.add_paragraph()
    r = intro.add_run(
        "A practical AI resource created for security managers and supervisors "
        "who need clearer reports, stronger handovers, better client updates "
        "and more consistent operational communication."
    )
    style_run(r, size=11.5, color=INK)


def copy_content(source, target):
    skipped_title = False
    for paragraph in source.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if not skipped_title and paragraph.style.name == "Title":
            skipped_title = True
            continue
        if paragraph.style.name == "Heading 1":
            new_p = target.add_paragraph(text, style="Heading 1")
        elif paragraph.style.name == "Heading 2":
            new_p = target.add_paragraph(text, style="Heading 2")
        elif paragraph.style.name == "Heading 3":
            new_p = target.add_paragraph(text, style="Heading 3")
        else:
            new_p = target.add_paragraph()
            label = None
            for candidate in [
                "Purpose:",
                "When to Use It:",
                "Example Input:",
                "Example Output:",
                "Copy-and-Paste Prompt:",
                "Advanced Version:",
                "Security Manager Tips:",
                "Challenge:",
                "Use AI To:",
                "Outcome:",
            ]:
                if text.startswith(candidate):
                    label = candidate
                    break
            if label:
                label_run = new_p.add_run(label + " ")
                style_run(label_run, bold=True, color=GOLD_DARK)
                body_run = new_p.add_run(text[len(label) :].strip())
                style_run(body_run, color=INK)
            else:
                run = new_p.add_run(text)
                style_run(run, color=INK)


def build_document(item):
    source = Document(item["src"])
    doc = Document()
    apply_base_styles(doc)
    add_footer(doc.sections[0])
    add_cover(doc, item["title"], item["subtitle"], item["delivery"])

    content_section = doc.add_section(WD_SECTION.NEW_PAGE)
    content_section.top_margin = Inches(0.85)
    content_section.bottom_margin = Inches(0.85)
    content_section.left_margin = Inches(1)
    content_section.right_margin = Inches(1)
    content_section.header_distance = Inches(0.492)
    content_section.footer_distance = Inches(0.492)
    add_footer(content_section)
    add_intro(doc, item["title"])
    copy_content(source, doc)
    doc.core_properties.title = item["title"]
    doc.core_properties.subject = item["subtitle"]
    doc.core_properties.author = BRAND
    doc.core_properties.keywords = (
        "AI prompts for security managers, ChatGPT prompts for security managers, "
        "security management templates"
    )
    item["docx"].parent.mkdir(parents=True, exist_ok=True)
    doc.save(item["docx"])
    print(item["docx"])


if __name__ == "__main__":
    for item in SOURCES:
        build_document(item)
